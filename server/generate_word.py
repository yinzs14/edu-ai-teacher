#!/usr/bin/env python3
"""
AI备课助手 — Word生成器
支持：
1. 家长沟通话术 Word 下载
2. 针对性练习 Word（学生版 + 教师版）
输入：JSON格式数据（通过stdin，type字段区分）
输出：docx文件路径
"""

import sys
import json
import datetime
import platform
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==================== 字体兼容 ====================
def _detect_font():
    """检测操作系统并返回合适的中文字体"""
    sys_name = platform.system()
    if sys_name == 'Windows':
        return FONT_NAME
    elif sys_name == 'Darwin':
        return 'PingFang SC'
    else:
        return 'Noto Sans SC'  # Linux (Ubuntu) 通用中文字体

FONT_NAME = _detect_font()

# ==================== 样式常量 ====================
TITLE_COLOR = RGBColor(0x1F, 0x29, 0x37)
ORANGE = RGBColor(0xF9, 0x73, 0x16)
BLUE = RGBColor(0x3B, 0x82, 0xF6)
GREEN = RGBColor(0x10, 0xB9, 0x81)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_header_footer(doc, student_name, teacher_name, doc_type):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = f'{student_name} · {doc_type}  |  教师：{teacher_name}'
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.style.font.size = Pt(9)
        hp.style.font.color.rgb = GRAY
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.font.name = FONT_NAME

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        today = datetime.date.today().strftime('%Y.%m.%d')
        fp.text = f'AI备课助手 · {today}'
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.font.name = FONT_NAME


# ==================== 家长沟通话术 Word ====================

def generate_communication_script(data, output_path):
    """生成家长沟通话术 Word 文档"""
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    student = data.get('studentName', '同学')
    teacher = data.get('teacherName', '老师')
    script = data.get('communicationScript', {})
    
    add_page_header_footer(doc, student, teacher, '家长沟通话术')

    # 标题
    title = doc.add_heading(f'{student} 家长沟通话术', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = TITLE_COLOR
        run.font.name = FONT_NAME

    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f'教师：{teacher}  |  日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    sr.font.size = Pt(11)
    sr.font.color.rgb = GRAY
    sr.font.name = FONT_NAME

    doc.add_paragraph()  # 空行

    sections = [
        ('📖 本阶段应掌握的知识', script.get('stageKnowledge', ''), '一、本阶段知识'),
        ('✅ 孩子已掌握的部分', script.get('mastered', ''), '二、已掌握部分'),
        ('⚠️ 有待提升的部分', script.get('weaknesses', ''), '三、有待提升'),
        ('💡 解决建议', script.get('solutions', ''), '四、解决建议'),
        ('💬 沟通要点提示', script.get('talkingTips', ''), '五、沟通要点'),
    ]

    for icon_title, content, formal_title in sections:
        if content:
            # 标题
            h = doc.add_heading(formal_title, level=2)
            for run in h.runs:
                run.font.color.rgb = ORANGE
                run.font.name = FONT_NAME
            
            # 内容
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.8
            p.paragraph_format.first_line_indent = Cm(0.7)
            r = p.add_run(content)
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            r.font.name = FONT_NAME
            
            doc.add_paragraph()  # 空行

    # 底部说明
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run('— 以上内容由 AI备课助手自动生成，仅供教师参考 —')
    nr.font.size = Pt(9)
    nr.font.color.rgb = GRAY
    nr.font.name = FONT_NAME

    doc.save(output_path)
    return output_path


# ==================== 针对性练习 Word ====================

def generate_practice_word(data, output_path, mode='student'):
    """
    生成针对性练习 Word 文档
    mode: 'student' 学生版（题在前，答案在最后）
          'teacher' 教师版（答案紧接题目，含知识点和解析）
    """
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    student = data.get('studentName', '同学')
    teacher = data.get('teacherName', '老师')
    exercises = data.get('exercises', [])
    weak_point = data.get('weakPoint', {})
    wp_name = weak_point.get('name', '综合练习')
    
    mode_label = '学生版' if mode == 'student' else '教师版（含解析）'
    add_page_header_footer(doc, student, teacher, f'针对性练习-{wp_name}')

    # ===== 封面 =====
    title = doc.add_heading(f'{wp_name} 针对性练习', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = TITLE_COLOR
        run.font.name = FONT_NAME

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f'{mode_label}  |  教师：{teacher}  |  {datetime.date.today().strftime("%Y.%m.%d")}')
    sr.font.size = Pt(11)
    sr.font.color.rgb = GRAY
    sr.font.name = FONT_NAME

    # 学生信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(f'姓名：________  日期：________  用时：________')
    ir.font.size = Pt(11)
    ir.font.color.rgb = GRAY
    ir.font.name = FONT_NAME

    doc.add_page_break()

    if mode == 'student':
        # ===== 学生版：题目在前 =====
        h = doc.add_heading(f'一、练习题（共 {len(exercises)} 题）', level=2)
        for run in h.runs:
            run.font.color.rgb = BLUE
            run.font.name = FONT_NAME

        for i, ex in enumerate(exercises):
            # 题号
            q_title = doc.add_paragraph()
            q_title.paragraph_format.space_before = Pt(12)
            qr = q_title.add_run(f'{i + 1}. {ex.get("title", "")}')
            qr.font.size = Pt(13)
            qr.font.bold = True
            qr.font.color.rgb = TITLE_COLOR
            qr.font.name = FONT_NAME

            # 题干
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Cm(0.5)
            q.paragraph_format.line_spacing = 1.8
            qr2 = q.add_run(ex.get('content', ''))
            qr2.font.size = Pt(12)
            qr2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            qr2.font.name = FONT_NAME

            # 答题区域
            a = doc.add_paragraph()
            a.paragraph_format.left_indent = Cm(0.5)
            ar = a.add_run('答：________________________________________')
            ar.font.size = Pt(12)
            ar.font.color.rgb = GRAY
            ar.font.name = FONT_NAME
            
            # 留白区域
            for _ in range(3):
                blank = doc.add_paragraph()
                blank.paragraph_format.left_indent = Cm(0.5)
                br = blank.add_run('')
                br.font.size = Pt(12)
            
            doc.add_paragraph()

        # ===== 答案在最后 =====
        doc.add_page_break()
        h_ans = doc.add_heading('二、参考答案', level=2)
        for run in h_ans.runs:
            run.font.color.rgb = GREEN
            run.font.name = FONT_NAME

        for i, ex in enumerate(exercises):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f'{i + 1}. {ex.get("answer", "略")}')
            r.font.size = Pt(12)
            r.font.color.rgb = GREEN
            r.font.name = FONT_NAME

    else:
        # ===== 教师版：答案紧接题目 =====
        h = doc.add_heading(f'针对性练习（教师版，共 {len(exercises)} 题）', level=2)
        for run in h.runs:
            run.font.color.rgb = BLUE
            run.font.name = FONT_NAME

        for i, ex in enumerate(exercises):
            # 题号与题干
            q_title = doc.add_paragraph()
            q_title.paragraph_format.space_before = Pt(14)
            qr = q_title.add_run(f'{i + 1}. {ex.get("title", "")}')
            qr.font.size = Pt(14)
            qr.font.bold = True
            qr.font.color.rgb = TITLE_COLOR
            qr.font.name = FONT_NAME

            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Cm(0.5)
            q.paragraph_format.line_spacing = 1.8
            qr2 = q.add_run(ex.get('content', ''))
            qr2.font.size = Pt(12)
            qr2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            qr2.font.name = FONT_NAME

            # 知识点标签
            tags = ex.get('tags', [])
            if tags:
                tag_p = doc.add_paragraph()
                tag_p.paragraph_format.left_indent = Cm(0.5)
                tag_r = tag_p.add_run('考查知识点：' + '、'.join(tags))
                tag_r.font.size = Pt(10)
                tag_r.font.color.rgb = BLUE
                tag_r.font.name = FONT_NAME

            # 答案
            ans_p = doc.add_paragraph()
            ans_p.paragraph_format.left_indent = Cm(0.5)
            ans_p.paragraph_format.space_before = Pt(4)
            ans_r = ans_p.add_run(f'答案：{ex.get("answer", "略")}')
            ans_r.font.size = Pt(12)
            ans_r.font.color.rgb = GREEN
            ans_r.font.name = FONT_NAME

            # 解析
            analysis = ex.get('analysis', '')
            if analysis:
                ana_p = doc.add_paragraph()
                ana_p.paragraph_format.left_indent = Cm(0.5)
                ana_p.paragraph_format.line_spacing = 1.6
                ana_r = ana_p.add_run(f'详细解析：{analysis}')
                ana_r.font.size = Pt(11)
                ana_r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                ana_r.font.name = FONT_NAME

            doc.add_paragraph()

    # 底部说明
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run('— AI备课助手 · 针对性练习自动生成 —')
    nr.font.size = Pt(9)
    nr.font.color.rgb = GRAY
    nr.font.name = FONT_NAME

    doc.save(output_path)
    return output_path


# ==================== 主入口 ====================

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python generate_word.py <output_path>', file=sys.stderr)
        sys.exit(1)

    raw = sys.stdin.read()
    if not raw or not raw.strip():
        print('Error: empty input data', file=sys.stderr)
        sys.exit(1)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f'Error: invalid JSON input: {e}', file=sys.stderr)
        sys.exit(1)

    output = sys.argv[1]
    
    doc_type = data.get('type', 'script')  # 'script' | 'practice_student' | 'practice_teacher'
    
    try:
        if doc_type == 'script':
            if not data.get('communicationScript'):
                print('Warning: no communicationScript data, generating empty document', file=sys.stderr)
            result = generate_communication_script(data, output)
        elif doc_type in ('practice_student', 'practice_teacher'):
            if not data.get('exercises'):
                print('Warning: no exercises data, generating empty document', file=sys.stderr)
            mode = 'student' if doc_type == 'practice_student' else 'teacher'
            result = generate_practice_word(data, output, mode)
        else:
            print(f'Unknown doc_type: {doc_type}', file=sys.stderr)
            sys.exit(1)

        print(result)
    except Exception as e:
        print(f'Error generating document: {e}', file=sys.stderr)
        sys.exit(1)
